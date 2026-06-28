#!/usr/bin/env python3
"""
csv_to_word.py
==================================================================
Convert a quiz CSV (sn, question, answer, category) into a clean,
reviewer-friendly Word (.docx) document.

USAGE
  python csv_to_word.py acts_quiz_1_14.csv
  python csv_to_word.py acts_quiz_1_14.csv --out acts_review.docx --title "Acts 1-14 Quiz"

The document groups questions by category, numbers them within each
category, and lays out each question with its answer underneath so a
reviewer can read and mark it up easily. Category codes are given
friendly headings (General, Spelling, etc.).
==================================================================
"""
import argparse
import csv
import os
import re
import sys
from collections import OrderedDict

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    sys.exit("This script needs python-docx. Install it with:\n"
             "  pip install python-docx --break-system-packages")

# Friendly headings for known category codes (same spirit as the build script)
CATEGORY_LABELS = {
    "general": "General Questions",
    "recitation": "Recitation",
    "recite": "Recitation",
    "fill": "Fill in the Gap",
    "fill-in-the-gap": "Fill in the Gap",
    "tof": "True or False",
    "true-or-false": "True or False",
    "tf": "True or False",
    "spell": "Spelling",
    "spelling": "Spelling",
    "golden": "Spelling",
    "golden-rush": "Spelling",
    "whosaid": "Who Said It?",
    "chapter": "Chapter Locator",
}

# On-screen / on-page order for known categories
CATEGORY_ORDER = {
    "General Questions": 1, "Recitation": 2, "Fill in the Gap": 3,
    "True or False": 4, "Spelling": 5, "Who Said It?": 6, "Chapter Locator": 7,
}

# Brand colors (Okota Baptist Church seal)
NAVY = RGBColor(0x0D, 0x0D, 0x0F)
RED = RGBColor(0xC8, 0x24, 0x2B)
GOLD = RGBColor(0xA8, 0x84, 0x1E)   # a slightly darker gold for print legibility
GREY = RGBColor(0x55, 0x55, 0x55)


def norm(s):
    return re.sub(r"\s+", " ", str(s)).strip()


def label_for(code):
    k = norm(code).lower().replace(" ", "-")
    return CATEGORY_LABELS.get(k, norm(code).title())


def read_rows(path):
    with open(path, newline="", encoding="utf-8-sig") as f:
        reader = csv.DictReader(f)
        reader.fieldnames = [norm(h).lower() for h in (reader.fieldnames or [])]
        rows = []
        for r in reader:
            d = {(k or ""): (v if v is not None else "") for k, v in r.items()}
            if any(norm(v) for v in d.values()):
                rows.append(d)
        return rows


def pick(row, *names):
    for n in names:
        if n in row and norm(row[n]):
            return norm(row[n])
    return ""


def set_cell_bg(cell, hex_color):
    """Shade a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_margins(cell, top=60, bottom=60, left=120, right=120):
    tcPr = cell._tc.get_or_add_tcPr()
    m = OxmlElement("w:tcMar")
    # schema order: top, start, bottom, end
    for tag, val in (("top", top), ("start", left), ("bottom", bottom), ("end", right)):
        node = OxmlElement(f"w:{tag}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        m.append(node)
    tcPr.append(m)


def add_page_number_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Page ")
    run.font.size = Pt(8)
    run.font.color.rgb = GREY
    # PAGE field
    fld1 = OxmlElement("w:fldSimple")
    fld1.set(qn("w:instr"), "PAGE")
    p._p.append(fld1)


def build_doc(rows, title, version, church, source_name):
    doc = Document()

    # Base style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    add_page_number_footer(section)

    # ---- Title block ----
    eyebrow = doc.add_paragraph()
    r = eyebrow.add_run((church or "").upper())
    r.font.size = Pt(9)
    r.font.bold = True
    r.font.color.rgb = GOLD
    eyebrow.space_after = Pt(0)

    h = doc.add_paragraph()
    rh = h.add_run(title)
    rh.font.size = Pt(22)
    rh.font.bold = True
    rh.font.color.rgb = NAVY
    h.space_after = Pt(2)

    sub = doc.add_paragraph()
    rs = sub.add_run(f"{version} · Review copy")
    rs.font.size = Pt(10)
    rs.font.color.rgb = GREY
    rs.italic = True

    # thin rule under the title
    rule = doc.add_paragraph()
    pPr = rule._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "C8242B")
    pbdr.append(bottom)
    pPr.append(pbdr)

    # ---- Group rows by category ----
    buckets = OrderedDict()
    for row in rows:
        cat = pick(row, "category", "cat", "type")
        q = pick(row, "question", "q", "prompt")
        a = pick(row, "answer", "a", "ans")
        if not (cat and q):
            continue
        label = label_for(cat)
        buckets.setdefault(label, []).append((q, a))

    ordered = sorted(buckets.items(),
                     key=lambda kv: (CATEGORY_ORDER.get(kv[0], 99), kv[0]))

    total = sum(len(v) for v in buckets.values())

    # ---- Contents summary ----
    csum = doc.add_paragraph()
    csum.paragraph_format.space_before = Pt(8)
    rc = csum.add_run("Contents:  ")
    rc.bold = True
    rc.font.size = Pt(10)
    parts = [f"{label} ({len(items)})" for label, items in ordered]
    rc2 = csum.add_run("   ".join(parts))
    rc2.font.size = Pt(10)
    rc2.font.color.rgb = GREY
    tot = doc.add_paragraph()
    rt = tot.add_run(f"Total questions: {total}")
    rt.font.size = Pt(10)
    rt.bold = True

    # ---- Each category ----
    for label, items in ordered:
        # Category heading
        ch = doc.add_paragraph()
        ch.paragraph_format.space_before = Pt(16)
        ch.paragraph_format.space_after = Pt(6)
        rch = ch.add_run(f"{label}  ")
        rch.font.size = Pt(15)
        rch.font.bold = True
        rch.font.color.rgb = NAVY
        rcount = ch.add_run(f"({len(items)})")
        rcount.font.size = Pt(11)
        rcount.font.color.rgb = GOLD
        rcount.bold = True
        # keep heading with next content
        ch.paragraph_format.keep_with_next = True

        # Two-column table: number | question + answer
        table = doc.add_table(rows=0, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.autofit = False
        table.allow_autofit = False
        # set column widths via grid
        widths = (Inches(0.35), Inches(6.4))
        # Force the grid column widths at the XML level (python-docx cell.width
        # alone is unreliable in some renderers).
        grid = table._tbl.find(qn("w:tblGrid"))
        if grid is not None:
            for gc, w in zip(grid.findall(qn("w:gridCol")), widths):
                gc.set(qn("w:w"), str(int(w.twips)))
        for i, item in enumerate(items, start=1):
            q, a = item
            row_cells = table.add_row().cells
            # number cell
            nc = row_cells[0]
            set_cell_margins(nc)
            np = nc.paragraphs[0]
            nr = np.add_run(str(i))
            nr.bold = True
            nr.font.size = Pt(11)
            nr.font.color.rgb = RED
            # content cell
            cc = row_cells[1]
            set_cell_margins(cc)
            qp = cc.paragraphs[0]
            qp.space_after = Pt(2)
            qr = qp.add_run(q)
            qr.font.size = Pt(11)
            # answer line
            ap = cc.add_paragraph()
            ap.paragraph_format.space_before = Pt(1)
            al = ap.add_run("Answer:  ")
            al.bold = True
            al.font.size = Pt(10)
            al.font.color.rgb = GOLD
            av = ap.add_run(a if a else "(no answer provided)")
            av.font.size = Pt(10)
            av.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            # set widths on cells
            nc.width = widths[0]
            cc.width = widths[1]
        # also set grid columns
        for col, w in zip(table.columns, widths):
            for cell in col.cells:
                cell.width = w

    return doc


def fix_settings_zoom(path):
    """python-docx writes a <w:zoom/> without the required percent attribute,
    which trips strict schema validation. Patch it so the file validates clean."""
    import zipfile, os, re
    with zipfile.ZipFile(path) as z:
        names = z.namelist()
        if "word/settings.xml" not in names:
            return
        settings = z.read("word/settings.xml").decode("utf-8")
    if "<w:zoom" not in settings or 'w:percent' in settings.split("<w:zoom", 1)[1][:60]:
        return
    new = re.sub(r'<w:zoom(?![^>]*w:percent)([^>]*?)(/?)>',
                 r'<w:zoom w:percent="100"\1\2>', settings)
    tmp = path + ".tmp"
    with zipfile.ZipFile(path) as zin, zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.namelist():
            data = zin.read(item)
            if item == "word/settings.xml":
                data = new.encode("utf-8")
            zout.writestr(item, data)
    os.replace(tmp, path)


def main():
    ap = argparse.ArgumentParser(description="Convert a quiz CSV to a Word review doc.")
    ap.add_argument("input", help="Path to the quiz .csv")
    ap.add_argument("--out", default=None, help="Output .docx (default: <input>_review.docx)")
    ap.add_argument("--title", default=None, help="Document title (default from filename)")
    ap.add_argument("--version", default="NIV", help="Bible version label (default: NIV)")
    ap.add_argument("--church", default="Okota Baptist Church", help="Church name")
    args = ap.parse_args()

    if not os.path.exists(args.input):
        sys.exit(f"File not found: {args.input}")

    rows = read_rows(args.input)
    if not rows:
        sys.exit("No rows found in the CSV.")

    title = args.title
    if not title:
        base = os.path.splitext(os.path.basename(args.input))[0]
        title = re.sub(r"[_\-]+", " ", base).replace("quiz", "Quiz").strip().title()
        title = title + ": Questions" if "quiz" not in title.lower() else title

    out = args.out or (os.path.splitext(args.input)[0] + "_review.docx")

    doc = build_doc(rows, title=title, version=args.version,
                    church=args.church, source_name=os.path.basename(args.input))
    doc.save(out)
    fix_settings_zoom(out)

    total = len(rows)
    print(f"\u2713 Wrote {out}")
    print(f"  Title: {title}")
    print(f"  Questions: {total}")


if __name__ == "__main__":
    main()